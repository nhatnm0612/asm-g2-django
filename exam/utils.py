from azure.storage.blob import BlobServiceClient, ContentSettings
from decouple import config
from pathlib import Path
from docx import Document
import re
import requests
from decouple import config
import json, os
from pathlib import WindowsPath
from uuid import uuid4

exam_info = {}
question_lst = []
blob_service_client = BlobServiceClient.from_connection_string(config("AZURE_CONNECTION_STRING"))
blob_paths = []

def search_value(pattern, text):
    match = re.search(pattern, text)
    if match:
         return match.group()
    else:
        return None

def push_image_to_blob(image_path: str, container_name: str ,blob_name: str):
    """Push image to blob

    Args:
        image_path (WindowsPath): _description_
        container_name (str): _description_
        blob_name (str): _description_
    """
    container_client = blob_service_client.get_container_client(container_name)  
    blob_client = container_client.get_blob_client(blob_name)
    content_settings = ContentSettings(content_type="image/png", override=True)

    with open(image_path, "rb") as data:
        blob_client.upload_blob(data, content_settings=content_settings, overwrite=True) 

def extract_docx_data(container_name: str, blob_name: str) -> tuple[list[list[str]], dict]:
    # Connect to Azure Storage
    container_client = blob_service_client.get_container_client(container_name)
    blob_client = container_client.get_blob_client(blob_name)

    # Save folder
    save_blob_folder = Path.home() / f"OneDrive/Documents/docx-processing"
    local_path = save_blob_folder / blob_name

    # Open blob_name in local folder
    with open(local_path, "wb+") as file:
        download_stream = blob_client.download_blob()
        file.write(download_stream.readall())

        # Save image in word file
        document = Document(local_path)
        
        # Read words file, create image in local documents and push back to blob
        for index, shape in enumerate(document.part.related_parts.values()):
            if hasattr(shape, "image"):
                image_bytes = shape.image.blob
                image_name = f"{uuid4()}.png"
                with open(save_blob_folder / image_name, "wb") as f:
                    f.write(image_bytes)
                    blob_paths.append(image_name)
                    # Push image extract from words file to container
                    push_image_to_blob(save_blob_folder / image_name, container_name, image_name)
        print(blob_paths)
        paragraphs = document.paragraphs
        headers_text = '\n'.join([paragraph.text for paragraph in paragraphs])

        #Init pattern
        subject_pattern, quiz_count_pattern, lecturer_pattern, date_pattern = r"Subject:\s*(.*)", r"Number of Quiz:\s*(\d+)", r"Lecturer:\s*(.*)", r"Date:\s*(.*)"

        # Extract headers of docx
        exam_info['subject'] = search_value(subject_pattern, headers_text).split(":")[-1].strip()
        exam_info['quiz_count'] = search_value(quiz_count_pattern, headers_text).split(":")[-1].strip()
        exam_info['lecturer'] = search_value(lecturer_pattern, headers_text).split(":")[-1].strip()
        exam_info['date'] = search_value(date_pattern, headers_text).split(":")[-1].strip()

        # Extract tables of docx
        tables = document.tables
        current_question = []
        for index,table in enumerate(tables):
            for row in table.rows:
                row_data = [cell.text for cell in row.cells]
                if row_data[0].startswith("QN"):
                    current_question = [row_data[1]]
                else:
                    current_question.append(row_data[1])
            current_question.append(blob_paths[index])
            question_lst.append(current_question)
        print(question_lst)

    # Clear env
    clear_folder = os.listdir(save_blob_folder)
    for file in clear_folder:
        os.remove(os.path.join(save_blob_folder, file))
    
    return question_lst, exam_info

def ask(content: str): 
    """Connect to GPT 3.5

    Args:
        content (str): _description_

    Returns:
        _type_: _description_
    """
    url = config("OPENAI_URL")
    
    querystring = {"api-version":"2024-04-01-preview"}
    
    payload = {"messages": [{"role": "user", "content": f"{content}"}], "temperature" : 1}

    headers = {
        'content-type': "application/json",
        'api-key': f"{config("OPENAI_KEY")}",
        'cache-control': "no-cache",
        'postman-token': "1b99988d-0914-6f27-a043-5421932e8f5b"
        }
    
    response = requests.post(url, json=payload, headers=headers, params=querystring)
    response_data = json.loads(response.text)
    return_content = response_data["choices"][0]["message"]["content"]
    return return_content

if __name__ == "__main__":
    extract_docx_data("asm-g2-docx-container", "Template.docx")